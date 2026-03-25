from __future__ import annotations

import base64
import csv
import io
import hashlib
import json
import logging
import mimetypes
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

try:
    import fitz
except Exception:  # pragma: no cover - optional dependency at runtime
    fitz = None

try:
    import pdfplumber
except Exception:  # pragma: no cover - optional dependency at runtime
    pdfplumber = None

try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency at runtime
    Document = None

try:
    import xlrd
except Exception:  # pragma: no cover - optional dependency at runtime
    xlrd = None

try:
    from openpyxl import Workbook, load_workbook
except Exception:  # pragma: no cover - optional dependency at runtime
    Workbook = None
    load_workbook = None
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, PlainTextResponse, Response, StreamingResponse
from openai import OpenAI
from sqlmodel import select
from dotenv import load_dotenv

from .chunking import chunk_text
from .metadata_extract import evidence_chunk_metadata
from .db import get_session, init_db
from .models import ChatSessionMessage, Project, ProjectDocument
from .schemas import (
    ChatMessageCreate,
    ChatMessageRead,
    ChatStreamMode,
    DocumentExportRequest,
    EvidenceDeleteRequest,
    OcrRebuildRequest,
    GenerateRequest,
    ModelProbeRequest,
    ProjectChatStreamRequest,
    ProjectCreate,
    ProjectDocumentUpdate,
    ProjectRead,
    ProjectStateRead,
    ProjectUpdate,
    RagQueryRequest,
    RagQueryResponse,
)
from .vector_store import (
    build_metadata_where,
    delete_project_collection,
    embed_texts,
    ensure_project_collection,
    query_project_chunks,
    upsert_project_chunks,
)

load_dotenv(Path(__file__).resolve().parent / ".env")

app_logger = logging.getLogger(__name__)

app = FastAPI(title="LawLaw Backend", version="0.2.8")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_origin_regex=r"^https?://(localhost|127\.0\.0\.1)(:\d+)?$",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

ALLOWED_EXTENSIONS = {
    ".zip",
    ".doc",
    ".docx",
    ".pdf",
    ".png",
    ".jpg",
    ".jpeg",
    ".txt",
    ".md",
    ".xlsx",
    ".xls",
    ".csv",
}
TEMPLATE_ALLOWED_EXTENSIONS = {".doc", ".docx", ".pdf"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
TEXT_EXTENSIONS = {".txt", ".md", ".doc", ".docx", ".pdf", ".xlsx", ".csv"}
BASE_UPLOAD_DIR = Path(tempfile.gettempdir()) / "lawlaw_uploads"
DOUBAO_BASE_URL = "https://ark.cn-beijing.volces.com/api/v3"
MAX_TEXT_CHARS_PER_FILE = 8000
CHAT_DOCUMENT_MARKDOWN_MAX = int(os.environ.get("CHAT_DOCUMENT_MARKDOWN_MAX", "120000"))


def project_dir_by_id(project_id: int) -> Path:
    return BASE_UPLOAD_DIR / str(project_id)

def _resolve_project_file_or_400(project_id: int, rel_path: str) -> Path:
    raw = (rel_path or "").strip().replace("\\", "/")
    if not raw or raw.startswith("/"):
        raise HTTPException(status_code=400, detail="invalid rel_path")
    if "\x00" in raw:
        raise HTTPException(status_code=400, detail="invalid rel_path")
    project_dir = project_dir_by_id(project_id).resolve()
    full = (project_dir / raw).resolve()
    if full == project_dir or project_dir not in full.parents:
        raise HTTPException(status_code=400, detail="invalid rel_path")
    if not full.exists() or not full.is_file():
        raise HTTPException(status_code=404, detail="file not found")
    if _is_junk_evidence_file(full):
        raise HTTPException(status_code=404, detail="file not found")
    if "template" in full.parts:
        raise HTTPException(status_code=404, detail="file not found")
    return full


def detect_file_type(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".zip":
        return "ZIP"
    if ext == ".pdf":
        return "PDF"
    if ext in {".doc", ".docx"}:
        return "WORD"
    if ext == ".xlsx":
        return "Excel"
    if ext == ".csv":
        return "CSV"
    if ext in IMAGE_EXTENSIONS:
        return "图片"
    if ext in {".txt", ".md"}:
        return "文本"
    return "其他"

def evidence_readability(path: Path) -> tuple[bool, str]:
    ext = path.suffix.lower()
    if _is_junk_evidence_file(path):
        return False, "系统文件/垃圾文件，不会被读取"
    if ext in IMAGE_EXTENSIONS:
        return True, ""
    if ext in {".pdf", ".txt", ".md", ".docx", ".xlsx", ".csv"}:
        return True, ""
    if ext == ".xls":
        return (
            False,
            ".xls 无法自动转为 .xlsx（已尝试 Python 与 LibreOffice；可能已加密或损坏），请另存为 .xlsx 或上传 .csv",
        )
    if ext == ".doc":
        return False, "暂不支持 .doc（建议转 .docx 或 PDF）"
    return False, "当前类型暂不支持读取"


def save_upload_file(uploaded: UploadFile, target_path: Path) -> None:
    with target_path.open("wb") as target:
        shutil.copyfileobj(uploaded.file, target)


def list_project_files(project_dir: Path) -> list[dict[str, str]]:
    files: list[dict[str, str]] = []
    name_counter: dict[str, int] = {}
    for item in sorted(project_dir.rglob("*")):
        if not item.is_file():
            continue
        # 不要把内部缓存目录当作“证据文件”展示
        if ".ocr_cache" in item.parts:
            continue
        # 兜底：即便缓存文件被放在其它目录，也不要展示
        lowered = item.name.lower()
        if lowered.endswith(".ocr.json") or lowered.endswith(".keyfacts.json"):
            continue
        if "template" in item.parts:
            continue
        # windows junk / zip junk that may slip in
        if item.name in {"desktop.ini"}:
            continue
        flat_name = item.name
        if flat_name in name_counter:
            name_counter[flat_name] += 1
            stem = Path(flat_name).stem
            suffix = Path(flat_name).suffix
            flat_name = f"{stem}_{name_counter[item.name]}{suffix}"
        else:
            name_counter[flat_name] = 1
        rel_path = item.relative_to(project_dir).as_posix()
        readable, reason = evidence_readability(item)
        files.append(
            {
                "name": flat_name,
                "type": detect_file_type(item),
                "rel_path": rel_path,
                "readable": "true" if readable else "false",
                "unreadable_reason": reason,
            }
        )
    return files


def trim_text(text: str, max_len: int = MAX_TEXT_CHARS_PER_FILE) -> str:
    cleaned = (text or "").strip()
    if len(cleaned) <= max_len:
        return cleaned
    return f"{cleaned[:max_len]}\n...(内容过长，已截断)"


def extract_pdf_text(pdf_path: Path) -> str:
    if pdfplumber is not None:
        try:
            pages: list[str] = []
            with pdfplumber.open(str(pdf_path)) as pdf:
                for page in pdf.pages:
                    pages.append(page.extract_text() or "")
            text = "\n".join(pages).strip()
            if text:
                return text
        except Exception:
            pass

    if fitz is not None:
        try:
            pages = []
            with fitz.open(str(pdf_path)) as doc:
                for page in doc:
                    pages.append(page.get_text("text") or "")
            return "\n".join(pages).strip()
        except Exception:
            return ""
    return ""


def extract_docx_text(docx_path: Path) -> str:
    if Document is None:
        return ""
    try:
        doc = Document(str(docx_path))
        return "\n".join(p.text for p in doc.paragraphs if p.text).strip()
    except Exception:
        return ""


def _sanitize_excel_sheet_title(name: str) -> str:
    raw = (name or "Sheet").strip() or "Sheet"
    bad = r':\/?*[]'
    cleaned = "".join("_" if c in bad else c for c in raw)
    return (cleaned[:31] if len(cleaned) > 31 else cleaned) or "Sheet"


def _unique_sheet_title(desired: str, used: set[str]) -> str:
    base = _sanitize_excel_sheet_title(desired)
    t = base
    n = 2
    while t in used:
        suf = f"_{n}"
        t = (base[: max(0, 31 - len(suf))] + suf) if len(base) + len(suf) > 31 else base + suf
        n += 1
    used.add(t)
    return t


def _find_libreoffice_soffice() -> str | None:
    """可执行文件路径；支持环境变量 LAWLAW_SOFFICE / LIBREOFFICE_SOFFICE。"""
    for key in ("LAWLAW_SOFFICE", "LIBREOFFICE_SOFFICE"):
        raw = (os.getenv(key) or "").strip()
        if raw and Path(raw).is_file():
            return raw
    for cmd in ("soffice", "libreoffice"):
        found = shutil.which(cmd)
        if found:
            return found
    mac = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    if mac.is_file():
        return str(mac)
    return None


def _convert_xls_via_libreoffice(xls_path: Path) -> Path | None:
    soffice = _find_libreoffice_soffice()
    if not soffice:
        return None
    out_path = _unique_xlsx_path(xls_path)
    tmp = Path(tempfile.mkdtemp(prefix="lawlaw_lo_", dir=xls_path.parent))
    try:
        proc = subprocess.run(
            [soffice, "--headless", "--convert-to", "xlsx", "--outdir", str(tmp), str(xls_path)],
            check=False,
            timeout=180,
            capture_output=True,
            text=True,
        )
        if proc.returncode != 0:
            app_logger.warning(
                "LibreOffice xls convert rc=%s stderr=%s file=%s",
                proc.returncode,
                (proc.stderr or "")[:500],
                xls_path.name,
            )
            return None
        produced = tmp / f"{xls_path.stem}.xlsx"
        if not produced.is_file():
            xlsx_list = list(tmp.glob("*.xlsx"))
            if not xlsx_list:
                app_logger.warning("LibreOffice produced no xlsx for %s", xls_path.name)
                return None
            produced = xlsx_list[0]
        shutil.move(str(produced), str(out_path))
        xls_path.unlink(missing_ok=True)
        app_logger.info("xls converted via LibreOffice -> %s", out_path.name)
        return out_path
    except subprocess.TimeoutExpired:
        app_logger.warning("LibreOffice convert timeout: %s", xls_path.name)
        return None
    except Exception as exc:
        app_logger.warning("LibreOffice convert failed %s: %s", xls_path.name, exc)
        return None
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def _unique_xlsx_path(for_xls: Path) -> Path:
    parent = for_xls.parent
    stem = for_xls.stem
    candidate = parent / f"{stem}.xlsx"
    if not candidate.exists():
        return candidate
    candidate = parent / f"{stem}_from_xls.xlsx"
    if not candidate.exists():
        return candidate
    n = 2
    while True:
        c = parent / f"{stem}_from_xls_{n}.xlsx"
        if not c.exists():
            return c
        n += 1


def _convert_xls_with_xlrd(xls_path: Path) -> Path | None:
    if xlrd is None or Workbook is None:
        return None
    book = None
    try:
        book = xlrd.open_workbook(str(xls_path), formatting_info=False, on_demand=True)
    except Exception as exc:
        app_logger.warning("xlrd open failed %s: %s", xls_path.name, exc)
        return None
    try:
        out_path = _unique_xlsx_path(xls_path)
        wb = Workbook()
        ws = wb.active
        titles_used: set[str] = set()
        for i in range(book.nsheets):
            xs = book.sheet_by_index(i)
            title = _unique_sheet_title(xs.name, titles_used)
            if i == 0:
                ws.title = title
            else:
                ws = wb.create_sheet(title=title)
            for row_idx in range(xs.nrows):
                ws.append([xs.cell_value(row_idx, c) for c in range(xs.ncols)])
        wb.save(str(out_path))
        xls_path.unlink(missing_ok=True)
        return out_path
    except Exception as exc:
        app_logger.warning("xls to xlsx convert failed %s: %s", xls_path.name, exc)
        return None
    finally:
        if book is not None:
            try:
                book.release_resources()
            except Exception:
                pass


def convert_xls_to_xlsx(xls_path: Path) -> Path | None:
    out = _convert_xls_with_xlrd(xls_path)
    if out is not None:
        return out
    if xls_path.is_file():
        return _convert_xls_via_libreoffice(xls_path)
    return None


def normalize_xls_files_under(project_dir: Path) -> None:
    for item in sorted(project_dir.rglob("*")):
        if not item.is_file() or item.suffix.lower() != ".xls":
            continue
        if ".ocr_cache" in item.parts or "template" in item.parts:
            continue
        if _is_junk_evidence_file(item):
            continue
        convert_xls_to_xlsx(item)


def extract_xlsx_text(xlsx_path: Path) -> str:
    if load_workbook is None:
        return ""
    try:
        wb = load_workbook(xlsx_path, read_only=True, data_only=True)
    except Exception:
        return ""
    try:
        parts: list[str] = []
        for ws in wb.worksheets:
            parts.append(f"## {ws.title}")
            for row in ws.iter_rows(values_only=True):
                line = "\t".join("" if v is None else str(v) for v in row)
                if line.strip():
                    parts.append(line)
            parts.append("")
        return "\n".join(parts).strip()
    except Exception:
        return ""
    finally:
        try:
            wb.close()
        except Exception:
            pass


def extract_csv_text(csv_path: Path) -> str:
    try:
        raw = csv_path.read_bytes()
    except Exception:
        return ""
    text: str | None = None
    for enc in ("utf-8-sig", "utf-8", "gb18030", "gbk", "latin-1"):
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        return ""
    try:
        sample = text[:4096] if len(text) > 4096 else text
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
        except csv.Error:
            dialect = csv.excel
        reader = csv.reader(io.StringIO(text), dialect)
        lines: list[str] = []
        for row in reader:
            lines.append("\t".join(cell or "" for cell in row))
        return "\n".join(lines).strip()
    except Exception:
        return ""


def extract_text_by_file(path: Path) -> str:
    ext = path.suffix.lower()
    # OCR 缓存优先（用于图片/扫描型 PDF 的预览与切块）
    try:
        rel = path.resolve().relative_to(BASE_UPLOAD_DIR.resolve())
        if rel.parts:
            project_dir = BASE_UPLOAD_DIR / rel.parts[0]
        else:
            project_dir = None
    except Exception:
        project_dir = None

    if project_dir and (ext in IMAGE_EXTENSIONS or ext == ".pdf"):
        try:
            fingerprint = compute_file_fingerprint(path)
            cache_file = _ocr_cache_file(project_dir, path)
            cached = _read_ocr_cache_text(cache_file, fingerprint)
            if cached is not None:
                return cached
        except Exception:
            # cache 读取异常不影响兜底提取
            pass

    if ext == ".pdf":
        return extract_pdf_text(path)
    if ext == ".docx":
        return extract_docx_text(path)
    if ext == ".xlsx":
        return extract_xlsx_text(path)
    if ext == ".csv":
        return extract_csv_text(path)
    if ext in {".txt", ".md"}:
        try:
            return path.read_text(encoding="utf-8", errors="ignore").strip()
        except Exception:
            return ""
    return ""


def load_template_text(project_dir: Path) -> str:
    template_dir = project_dir / "template"
    if not template_dir.exists():
        return "（未检测到参考模板）"
    files = [f for f in sorted(template_dir.iterdir()) if f.is_file()]
    if not files:
        return "（未检测到参考模板）"

    template_file = files[0]
    ext = template_file.suffix.lower()
    if ext == ".pdf":
        text = extract_pdf_text(template_file)
    elif ext == ".docx":
        text = extract_docx_text(template_file)
    else:
        text = "（.doc 文件暂不支持直接解析，请转为 .docx / .pdf）"

    text = trim_text(text)
    if not text:
        return f"（模板 {template_file.name} 未提取到可用文本）"
    return text


def encode_image_to_data_url(path: Path) -> str:
    mime_type, _ = mimetypes.guess_type(str(path))
    if not mime_type:
        mime_type = "image/jpeg"
    binary = path.read_bytes()
    encoded = base64.b64encode(binary).decode("utf-8")
    return f"data:{mime_type};base64,{encoded}"


def _compute_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def compute_file_fingerprint(path: Path) -> dict[str, Any]:
    st = path.stat()
    return {
        "sha256": _compute_sha256(path),
        "size": int(st.st_size),
        "mtime_ns": int(getattr(st, "st_mtime_ns", int(st.st_mtime * 1e9))),
    }


def _ocr_cache_file(project_dir: Path, file_path: Path) -> Path:
    cache_dir = project_dir / ".ocr_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    # 以文件名作为缓存文件名；内容变更依赖 fingerprint 校验
    return cache_dir / f"{file_path.name}.ocr.json"


def _read_ocr_cache_text(cache_file: Path, expected_fingerprint: dict[str, Any]) -> str | None:
    try:
        if not cache_file.exists():
            return None
        raw = cache_file.read_text(encoding="utf-8", errors="ignore").strip()
        if not raw:
            return None
        body = json.loads(raw)
        if body.get("fingerprint") == expected_fingerprint:
            text = body.get("text")
            if isinstance(text, str):
                return text
        return None
    except Exception:
        return None


def _write_ocr_cache_text(cache_file: Path, fingerprint: dict[str, Any], text: str) -> None:
    try:
        cache_file.write_text(
            json.dumps({"fingerprint": fingerprint, "text": text}, ensure_ascii=False),
            encoding="utf-8",
        )
    except Exception:
        # 缓存失败不影响 OCR 结果本身
        pass


KEYFACTS_PROMPT_VERSION = 2


def _keyfacts_cache_file(project_dir: Path, file_path: Path) -> Path:
    cache_dir = project_dir / ".ocr_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    return cache_dir / f"{file_path.name}.keyfacts.json"


def _read_keyfacts_cache(cache_file: Path, expected_fingerprint: dict[str, Any]) -> dict[str, str] | None:
    try:
        if not cache_file.exists():
            return None
        raw = cache_file.read_text(encoding="utf-8", errors="ignore").strip()
        if not raw:
            return None
        body = json.loads(raw)
        if body.get("fingerprint") != expected_fingerprint:
            return None
        # prompt 升级后，旧 keyfacts 可能不符合新字段要求，因此当作 missing 重新生成
        if body.get("keyfacts_version") != KEYFACTS_PROMPT_VERSION:
            return None
        full = body.get("keyfacts_full")
        compact = body.get("keyfacts_compact")
        if isinstance(full, str) and isinstance(compact, str):
            return {"keyfacts_full": full, "keyfacts_compact": compact}
        return None
    except Exception:
        return None


def _write_keyfacts_cache(
    cache_file: Path,
    fingerprint: dict[str, Any],
    keyfacts_full: str,
    keyfacts_compact: str,
) -> None:
    try:
        cache_file.write_text(
            json.dumps(
                {
                    "fingerprint": fingerprint,
                    "keyfacts_version": KEYFACTS_PROMPT_VERSION,
                    "keyfacts_full": keyfacts_full,
                    "keyfacts_compact": keyfacts_compact,
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
    except Exception:
        # 缓存失败不影响主流程
        pass


def _split_keyfacts_sections(raw: str) -> tuple[str, str]:
    full_marker = "<<<KEYFACTS_FULL>>>"
    compact_marker = "<<<KEYFACTS_COMPACT>>>"
    if full_marker in raw and compact_marker in raw:
        full_start = raw.index(full_marker) + len(full_marker)
        compact_start = raw.index(compact_marker)
        full = raw[full_start:compact_start].strip()
        compact = raw[compact_start + len(compact_marker) :].strip()
        return full, compact
    # fallback：没有分隔符时，把整段当作 full
    return raw.strip(), ""


def _guess_keyfacts_doc_type_guess(*, file_path: Path, source_text: str) -> str:
    """
    轻量级证据类型猜测：主要用于控制聊天类证据是否需要输出“事实经过/沟通时间线摘要”。
    """
    name_l = (file_path.name or "").lower()
    sample_l = (source_text or "")[:1200].lower()
    chat_keywords = [
        "微信",
        "qq",
        "聊天",
        "聊天记录",
        "对话",
        "短信",
        "通话记录",
        "通话",
        "语音",
        "截图",
        "撤回",
        "已读",
        "未读",
        "群聊",
        "消息",
    ]
    if any(k in name_l for k in chat_keywords) or any(k in sample_l for k in chat_keywords):
        return "chat_evidence"
    return "other"


def _keyfacts_extract_from_ocr_text(
    *,
    api_key: str,
    chat_vision_endpoint_id: str,
    ocr_text: str,
    file_name: str,
    doc_type_guess: str,
) -> tuple[str, str]:
    client = OpenAI(base_url=DOUBAO_BASE_URL, api_key=api_key)

    system_prompt = (
        "你是资深法律文书信息抽取器。输入是一段 OCR/文本。\n"
        "请只做“信息抽取”，不要进行解释、推理、也不要做法律结论。\n"
        "输出为纯文本，不要使用任何 code block，不要输出 JSON。\n"
        "必须严格遵守下面输出格式与约束：\n"
        "1) 输出分为两个部分：KEYFACTS_FULL 与 KEYFACTS_COMPACT。\n"
        "2) KEYFACTS_FULL：A-J（固定字段块）必须全部出现；如果某字段在文本中找不到，就写“未找到”（未找到字段不需要 snippet）。命中字段必须给值 + snippet。\n"
        "2.1) 如果 doc_type_guess == 'chat_evidence'，则还必须出现字段块 K（事实经过/沟通时间线摘要）；若找不到就写“未找到”（K 的未找到不需要 snippet）。\n"
        "3) KEYFACTS_COMPACT：只保留 KEYFACTS_FULL 中“找到了内容”的字段块；未命中字段块省略；不再输出 snippet。（但对聊天证据：字段块 K 即使未找到也必须出现，值写“未找到”。）\n"
        "2) 对法律条款/法条/章/节：必须使用原文的完整表达（不可缩写、不可省略号）。\n"
        "3) 对金额/法院/时间/地区/地址/案号：同样使用原文完整表达（不可缩写、不可省略）。\n"
        "4) 命中字段块必须包含：字段名 + 值 + snippet（snippet 必须是可核对的原文短片段；允许更长以不丢信息）。\n"
        "5) 在输出末尾，不要附加任何额外说明。\n"
        "\n"
        "输出必须包含两个部分，并使用固定分隔符：\n"
        "<<<KEYFACTS_FULL>>>\n"
        "（字段块 A-J 全部出现；未命中写“未找到”；命中带 snippet；聊天证据可能还会包含 K 字段块）\n"
        "<<<KEYFACTS_COMPACT>>>\n"
        "（压缩版：只保留命中字段名 + 值；不再输出 snippet）\n"
    )

    user_prompt = (
        "请根据下面“文件信息 + doc_type_guess + OCR/文本”抽取关键要素：\n\n"
        f"文件名：{file_name}\n"
        f"doc_type_guess：{doc_type_guess}\n\n"
        f"OCR/文本如下（可能很长）：\n{ocr_text}\n\n"
        "字段清单（A-J 固定字段块）：\n"
        "A. 文书类型与程序阶段\n"
        "B. 案号与法院信息（法院名称、分院/法庭如有）\n"
        "C. 时间信息（立案/受理/开庭/判决-裁定/送达/生效/执行立案与期限；去重全列出）\n"
        "D. 地区与地址类（住所或营业地、送达地址、联系地址等）\n"
        "E. 当事人信息（原告/被告/第三人、证件号或统一代码如有、代理人/律师/委托代理人如有）\n"
        "F. 金额与财产性请求（欠款本金/剩余本金、利息/利率与计算期间、违约金、赔偿/损失、诉讼费/执行费等）\n"
        "G. 证据/附件清单要点（证据名称；关键证据对应的金额/时间/主体要点）\n"
        "H. 争议焦点/请求事项（请求事项列表、争议焦点摘要）\n"
        "I. 条款/法条/章节（合同条款编号、法条条款如《民法典》第X条、适用法律、章/节）\n"
        "J. 关键数字（能定位到、且未来你会用来追溯的数字类别：标的编号、条款编号、案号数字片段等）\n"
        "\n"
        "聊天证据额外字段（仅当 doc_type_guess == 'chat_evidence' 时输出）：\n"
        "K. 事实经过/沟通时间线摘要（按时间点列关键对话/沟通事件/承诺/争议点；每个时间点尽量保留原文句式并附 snippet）\n"
    )

    completion = client.chat.completions.create(
        model=chat_vision_endpoint_id,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0,
        max_tokens=4096,
    )
    raw = completion.choices[0].message.content or ""
    full, compact = _split_keyfacts_sections(raw)
    return full, compact


def get_or_create_keyfacts_text(
    *,
    project_dir: Path,
    file_path: Path,
    api_key: str,
    chat_vision_endpoint_id: str,
    ocr_text: str,
    fingerprint: dict[str, Any],
    rekeyfacts_if_empty: bool = False,
    force: bool = False,
) -> dict[str, str]:
    cache_file = _keyfacts_cache_file(project_dir, file_path)
    if not force:
        cached = _read_keyfacts_cache(cache_file, fingerprint)
        if cached is not None:
            # 只有 keyfacts_full 非空时才直接复用；如果为空则继续生成
            full = str(cached.get("keyfacts_full") or "").strip()
            if full:
                return cached

    if not str(ocr_text).strip() or not chat_vision_endpoint_id:
        return {"keyfacts_full": "", "keyfacts_compact": ""}

    # KeyFacts 生成只依赖 OCR 文本；失败时回退为 missing（前端会提示先重扫）
    try:
        file_name = file_path.name
        doc_type_guess = _guess_keyfacts_doc_type_guess(file_path=file_path, source_text=ocr_text)
        keyfacts_full, keyfacts_compact = _keyfacts_extract_from_ocr_text(
            api_key=api_key,
            chat_vision_endpoint_id=chat_vision_endpoint_id,
            ocr_text=ocr_text,
            file_name=file_name,
            doc_type_guess=doc_type_guess,
        )
        keyfacts_full = str(keyfacts_full or "").strip()
        keyfacts_compact = str(keyfacts_compact or "").strip()

        _write_keyfacts_cache(
            cache_file=cache_file,
            fingerprint=fingerprint,
            keyfacts_full=keyfacts_full,
            keyfacts_compact=keyfacts_compact,
        )
        return {"keyfacts_full": keyfacts_full, "keyfacts_compact": keyfacts_compact}
    except Exception as exc:
        app_logger.error("KeyFacts 生成失败：%s（%s）", file_path.name, exc)
        return {"keyfacts_full": "", "keyfacts_compact": ""}


def _sanitize_ocr_output(text: str) -> str:
    t = (text or "").strip()
    if not t:
        return ""
    # 兜底：如果模型把内容包在 code block 里，尽量只取 code 内
    if "```" in t:
        m = re.search(r"```[a-zA-Z]*\s*([\s\S]*?)\s*```", t)
        if m:
            t = (m.group(1) or "").strip()
    return t


def _ocr_data_url_to_text(
    *,
    api_key: str,
    chat_vision_endpoint_id: str,
    image_data_url: str,
    max_tokens: int = 4096,
) -> str:
    client = OpenAI(base_url=DOUBAO_BASE_URL, api_key=api_key)
    system_prompt = (
        "你是高精度 OCR 引擎。请只识别图片中的文字并输出纯文本。\n"
        "不要输出任何解释、标题、JSON、代码块外壳或额外前后缀。\n"
        "如果无法识别某些内容可以忽略，但必须保持原有换行尽量合理。"
    )
    messages: list[dict[str, Any]] = [
        {"role": "system", "content": system_prompt},
        {
            "role": "user",
            "content": [
                {"type": "text", "text": "开始识别图片文字，只输出纯文本。"},
                {"type": "image_url", "image_url": {"url": image_data_url}},
            ],
        },
    ]
    completion = client.chat.completions.create(
        model=chat_vision_endpoint_id,
        messages=messages,
        temperature=0,
        max_tokens=max_tokens,
    )
    content = completion.choices[0].message.content or ""
    return _sanitize_ocr_output(content)


def get_or_create_ocr_text(
    *,
    project_dir: Path,
    file_path: Path,
    api_key: str,
    chat_vision_endpoint_id: str,
    ocr_dpi: int = 200,
    reocr_if_empty: bool = False,
) -> str:
    if not chat_vision_endpoint_id:
        # OCR 配置缺失时不要直接让上传/重建流程崩溃：
        # - 图片：当前没有 OCR，返回空（允许上层继续跳过）
        # - PDF：退回到已有文本层提取（保证兼容）
        ext = file_path.suffix.lower()
        if ext == ".pdf":
            return extract_text_by_file(file_path)
        return ""

    fingerprint = compute_file_fingerprint(file_path)
    cache_file = _ocr_cache_file(project_dir, file_path)
    cached = _read_ocr_cache_text(cache_file, fingerprint)
    if cached is not None:
        if reocr_if_empty and not str(cached).strip():
            cached = None
        else:
            # OCR 已命中但 keyfacts 缓存可能缺失/为空：只在 keyfacts_full 为空时补生成
            try:
                _ = get_or_create_keyfacts_text(
                    project_dir=project_dir,
                    file_path=file_path,
                    api_key=api_key,
                    chat_vision_endpoint_id=chat_vision_endpoint_id,
                    ocr_text=cached,
                    fingerprint=fingerprint,
                    force=False,
                )
            except Exception:
                pass
            return cached

    ext = file_path.suffix.lower()
    try:
        if ext in IMAGE_EXTENSIONS:
            data_url = encode_image_to_data_url(file_path)
            text = _ocr_data_url_to_text(
                api_key=api_key,
                chat_vision_endpoint_id=chat_vision_endpoint_id,
                image_data_url=data_url,
            )
        elif ext == ".pdf":
            if fitz is None:
                raise RuntimeError("缺少 fitz（PyMuPDF），无法渲染 PDF 页并进行 OCR")

            pages_text: list[str] = []
            with fitz.open(str(file_path)) as doc:
                mat = fitz.Matrix(ocr_dpi / 72, ocr_dpi / 72)
                for page_index in range(len(doc)):
                    page = doc.load_page(page_index)
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    img_bytes = pix.tobytes("jpg")
                    data_url = f"data:image/jpeg;base64,{base64.b64encode(img_bytes).decode('utf-8')}"
                    page_text = _ocr_data_url_to_text(
                        api_key=api_key,
                        chat_vision_endpoint_id=chat_vision_endpoint_id,
                        image_data_url=data_url,
                    )
                    if page_text:
                        pages_text.append(page_text)
            text = "\n\n".join(pages_text).strip()
        else:
            # OCR 只覆盖图片/PDF；其他类型由已有 extract_text_by_file 逻辑兜底
            text = extract_text_by_file(file_path)
    except Exception as exc:
        app_logger.error("OCR 失败：%s（%s）", file_path.name, exc)
        # OCR 失败时，尽量回退到“原有文本层”提取，避免把 cache 写成空导致降级。
        if ext == ".pdf":
            text = extract_pdf_text(file_path)
        elif ext in IMAGE_EXTENSIONS:
            text = ""
        else:
            text = extract_text_by_file(file_path)

    _write_ocr_cache_text(cache_file, fingerprint, text)

    # OCR 实际生成/重扫时，同步生成 KeyFacts（并缓存，避免每次打开都调模型）
    try:
        keyfacts = get_or_create_keyfacts_text(
            project_dir=project_dir,
            file_path=file_path,
            api_key=api_key,
            chat_vision_endpoint_id=chat_vision_endpoint_id,
            ocr_text=text,
            fingerprint=fingerprint,
            force=True,
        )
        # 仅为了避免“无内容也写入缓存”影响前端判断；KeyFacts 缓存函数本身已做空保护
        _ = keyfacts
    except Exception:
        # KeyFacts 不影响 OCR/向量入库主流程
        pass

    return text

def _is_junk_evidence_file(path: Path) -> bool:
    # macOS zip junk
    if "__MACOSX" in path.parts:
        return True
    name = path.name
    if name.startswith("._"):
        return True
    # Thumbs.db etc.
    if name in {".DS_Store", "Thumbs.db"}:
        return True
    # Some zip tools create zero-byte placeholders
    try:
        if path.stat().st_size == 0:
            return True
    except Exception:
        return True
    # Extremely small "images" are often AppleDouble or corrupt; skip in evidence flow.
    if path.suffix.lower() in IMAGE_EXTENSIONS:
        try:
            if path.stat().st_size < 1024:
                return True
        except Exception:
            return True
    return False


def _safe_zip_member_name(name: str) -> str:
    # Normalize separators, drop drive letters, prevent absolute paths.
    name = name.replace("\\", "/")
    while name.startswith("/"):
        name = name[1:]
    name = name.replace("\x00", "")
    return name


def _decode_zip_name(raw_name: str) -> str:
    # If Zip has UTF-8 flag, Python already decoded correctly; keep as-is.
    # Otherwise, many tools store raw bytes in CP437; re-decode best-effort.
    try:
        raw_bytes = raw_name.encode("cp437", errors="strict")
    except Exception:
        return raw_name
    for enc in ("utf-8", "gbk", "utf-16", "shift_jis"):
        try:
            decoded = raw_bytes.decode(enc)
            if decoded:
                return decoded
        except Exception:
            continue
    return raw_name


def extract_zip_safely(zip_path: Path, target_dir: Path) -> None:
    with zipfile.ZipFile(zip_path, "r") as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            name = _safe_zip_member_name(info.filename)
            name = _decode_zip_name(name)
            name = _safe_zip_member_name(name)
            if not name:
                continue
            # Skip macOS junk early
            if name.startswith("__MACOSX/") or "/__MACOSX/" in name or "/._" in name or name.startswith("._"):
                continue
            out_path = (target_dir / name).resolve()
            base = target_dir.resolve()
            if base not in out_path.parents and out_path != base:
                continue
            out_path.parent.mkdir(parents=True, exist_ok=True)
            with zf.open(info, "r") as src, out_path.open("wb") as dst:
                shutil.copyfileobj(src, dst)


def iter_evidence_files(project_dir: Path) -> Iterable[Path]:
    for item in sorted(project_dir.rglob("*")):
        if not item.is_file():
            continue
        # 排除内部缓存目录（OCR/KeyFacts 缓存不属于证据）
        if ".ocr_cache" in item.parts:
            continue
        lowered = item.name.lower()
        if lowered.endswith(".ocr.json") or lowered.endswith(".keyfacts.json"):
            continue
        readable, _ = evidence_readability(item)
        if not readable:
            continue
        if "template" in item.parts:
            continue
        if item.suffix.lower() == ".zip":
            continue
        yield item


def build_evidence_content_items(project_dir: Path) -> list[dict[str, Any]]:
    content: list[dict[str, Any]] = []
    for file_path in iter_evidence_files(project_dir):
        ext = file_path.suffix.lower()
        file_name = file_path.name
        if ext in IMAGE_EXTENSIONS:
            ocr_text = trim_text(extract_text_by_file(file_path))
            data_url = encode_image_to_data_url(file_path)
            if ocr_text:
                content.append({"type": "text", "text": f"[证据图片 doc:{file_name}]\n{ocr_text}"})
            else:
                content.append({"type": "text", "text": f"[证据图片 doc:{file_name}]"})
            content.append({"type": "image_url", "image_url": {"url": data_url}})
            continue
        if ext in TEXT_EXTENSIONS:
            text = trim_text(extract_text_by_file(file_path))
            if text:
                content.append({"type": "text", "text": f"[证据文本 doc:{file_name}]\n{text}"})
    if not content:
        content.append({"type": "text", "text": "（未检测到可读取的证据内容）"})
    return content


def _env_int(name: str, default: int) -> int:
    raw = (os.getenv(name) or "").strip()
    if not raw:
        return default
    try:
        return int(raw)
    except Exception:
        return default


def estimate_prompt_tokens(*, system_prompt: str, instruction: str, evidence_items: list[dict[str, Any]]) -> int:
    """
    粗估 token：不追求精确，只用于避免“超上下文导致生成失败”。
    中文按 ~3.2 字符/Token（偏保守）估算；图片 data_url 视为巨量开销，给一个高惩罚。
    """

    def est_text_tokens(text: str) -> int:
        t = (text or "").strip()
        if not t:
            return 0
        return max(1, (len(t) + 2) // 3)  # conservative-ish

    total = est_text_tokens(system_prompt) + est_text_tokens(instruction)
    for item in evidence_items:
        if not isinstance(item, dict):
            continue
        if item.get("type") == "text":
            total += est_text_tokens(str(item.get("text") or ""))
        elif item.get("type") == "image_url":
            # image_url 在 payload 里可能包含 base64 data_url；这里用高惩罚粗估
            total += 40_000
    return total


def build_evidence_content_items_compact(project_dir: Path) -> list[dict[str, Any]]:
    """
    压缩档：优先使用每个文件的 keyfacts_compact（字段名+值），避免上下文超限导致生成失败。
    - 不发送 image_url（避免 data_url 带来的巨大开销）
    - 若 keyfacts 缺失：回退 extract_text_by_file，并强截断到较小长度
    """

    content: list[dict[str, Any]] = []
    for file_path in iter_evidence_files(project_dir):
        file_name = file_path.name
        fingerprint = compute_file_fingerprint(file_path)
        cache_file = _keyfacts_cache_file(project_dir, file_path)
        cached = _read_keyfacts_cache(cache_file, fingerprint) or {}
        compact = str(cached.get("keyfacts_compact") or "").strip()
        if not compact:
            # 兜底：没有 KeyFacts 时，至少给一点原文（但要强截断，防止超限）
            raw = (extract_text_by_file(file_path) or "").strip()
            if raw:
                compact = trim_text(raw, max_len=3000)
        if not compact:
            continue
        content.append({"type": "text", "text": f"[证据要点 doc:{file_name}]\n{compact}"})
    if not content:
        content.append({"type": "text", "text": "（未检测到可读取的证据内容）"})
    return content


def sse_event(payload: dict[str, Any]) -> str:
    return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"


def _truncate_for_chat_context(text: str, max_chars: int = CHAT_DOCUMENT_MARKDOWN_MAX) -> str:
    t = text or ""
    if len(t) <= max_chars:
        return t
    return t[:max_chars] + "\n\n[... 后文已因长度截断，如需全文改写请分段处理 ...]"


def _chat_history_openai_messages(history: Any) -> list[dict[str, str]]:
    from .schemas import ProjectChatHistoryItem

    out: list[dict[str, str]] = []
    seq = history if isinstance(history, list) else []
    for item in seq[-20:]:
        if not isinstance(item, ProjectChatHistoryItem):
            continue
        role = (item.role or "").strip().lower()
        if role not in ("user", "assistant"):
            continue
        content = (item.content or "").strip()
        if not content:
            continue
        out.append({"role": role, "content": content})
    return out


def iter_rag_query_sse(payload: RagQueryRequest) -> Iterable[str]:
    query = payload.query.strip()
    api_key = resolve_api_key(payload.api_key)
    if not query:
        yield sse_event({"type": "error", "message": "query is required"})
        yield sse_event({"type": "done", "sources": []})
        return
    if not api_key:
        yield sse_event({"type": "error", "message": "缺少 API Key，请在齿轮填写或在 backend/.env 配置 ARK_API_KEY"})
        yield sse_event({"type": "done", "sources": []})
        return

    project_dir = project_dir_by_id(payload.project_id)
    if not project_dir.exists():
        yield sse_event({"type": "error", "message": "project data not found"})
        yield sse_event({"type": "done", "sources": []})
        return

    meta_where = build_metadata_where(
        filter_doc_type=payload.filter_doc_type or None,
        filter_source_file=payload.filter_source_file or None,
    )
    try:
        embed_id = require_embedding_id(
            payload.embedding_endpoint_id,
            payload.embedding_model_id,
        )
        chat_id = require_chat_vision_id(
            payload.chat_vision_endpoint_id,
            payload.chat_vision_model_id,
        )
    except HTTPException as exc:
        detail = exc.detail
        msg = detail if isinstance(detail, str) else str(detail)
        yield sse_event({"type": "error", "message": msg})
        yield sse_event({"type": "done", "sources": []})
        return

    try:
        retrieval = query_project_chunks(
            project_id=payload.project_id,
            query=query,
            top_k=payload.top_k,
            embedding_model_id=embed_id,
            where=meta_where,
            api_key=api_key,
        )
    except Exception as exc:
        yield sse_event({"type": "error", "message": f"RAG 检索失败：{exc}"})
        yield sse_event({"type": "done", "sources": []})
        return
    docs = retrieval.get("documents", [])
    metas = retrieval.get("metadatas", [])
    if not docs:
        yield sse_event(
            {
                "type": "delta",
                "content": "当前项目还没有可检索内容，请先上传并完成向量入库。",
            }
        )
        yield sse_event({"type": "done", "sources": []})
        return

    context_lines: list[str] = []
    source_names: list[str] = []
    for idx, doc in enumerate(docs):
        meta = metas[idx] if idx < len(metas) else {}
        doc_name = str(meta.get("source_file") or meta.get("doc_name", "未知文件"))
        source_names.append(doc_name)
        context_lines.append(f"[doc:{doc_name}] {doc}")
    context = "\n\n".join(context_lines)

    unique_sources: list[str] = []
    for name in source_names:
        if name not in unique_sources:
            unique_sources.append(name)

    system_prompt = (
        "你是资深非诉律师的项目检索助手。"
        "请仅基于给出的检索片段回答用户问题，不能编造事实。"
        "回答要简洁专业，并在关键句末用 [doc:文件名] 标注来源。"
    )
    user_prompt = f"用户问题：{query}\n\n检索片段：\n{context}"

    try:
        client = OpenAI(base_url=DOUBAO_BASE_URL, api_key=api_key)
        stream = client.chat.completions.create(
            model=chat_id,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            stream=True,
            temperature=0.2,
        )
        yield sse_event({"type": "sources", "sources": unique_sources})
        for chunk in stream:
            delta = ""
            if chunk.choices and chunk.choices[0].delta:
                delta = chunk.choices[0].delta.content or ""
            if delta:
                yield sse_event({"type": "delta", "content": delta})
        yield sse_event({"type": "done", "sources": unique_sources})
    except Exception as exc:
        yield sse_event({"type": "error", "message": str(exc)})


def upsert_project_evidence(
    project_id: int,
    chat_vision_endpoint_id: str | None = None,
    embedding_model_id: str | None = None,
    api_key: str | None = None,
) -> dict[str, Any]:
    project_dir = project_dir_by_id(project_id)
    if not project_dir.exists():
        return {"upserted": 0, "files": 0}

    total_upserted = 0
    processed_files = 0
    for file_path in iter_evidence_files(project_dir):
        # 上传/重建阶段：对图片与扫描 PDF 走 OCR（带缓存），其他类型兜底抽取文本
        text = get_or_create_ocr_text(
            project_dir=project_dir,
            file_path=file_path,
            api_key=api_key or "",
            chat_vision_endpoint_id=chat_vision_endpoint_id or "",
        )
        text = text or ""
        chunks = chunk_text(text)
        if not chunks:
            continue
        preview = text[:800] if text else ""
        chunk_meta = evidence_chunk_metadata(file_path, preview)
        result = upsert_project_chunks(
            project_id=project_id,
            doc_name=file_path.name,
            chunks=chunks,
            embedding_model_id=embedding_model_id,
            chunk_metadata=chunk_meta,
            api_key=api_key,
        )
        total_upserted += int(result["upserted"])
        processed_files += 1
    return {"upserted": total_upserted, "files": processed_files}


def get_project_or_404(project_id: int) -> Project:
    with get_session() as session:
        project = session.get(Project, project_id)
        if not project:
            raise HTTPException(status_code=404, detail="project not found")
        return project


def resolve_chat_vision_id(
    endpoint_id: str | None = None,
    legacy_model_id: str | None = None,
) -> str:
    """请求体 ep- 接入点优先，其次兼容旧字段，最后读 .env。"""
    return (
        (endpoint_id or "").strip()
        or (legacy_model_id or "").strip()
        or (os.getenv("CHAT_VISION_ENDPOINT_ID") or "").strip()
        or (os.getenv("CHAT_VISION_MODEL_ID") or "").strip()
    )


def resolve_embedding_id(
    endpoint_id: str | None = None,
    legacy_model_id: str | None = None,
) -> str:
    return (
        (endpoint_id or "").strip()
        or (legacy_model_id or "").strip()
        or (os.getenv("EMBEDDING_ENDPOINT_ID") or "").strip()
        or (os.getenv("EMBEDDING_MODEL_ID") or "").strip()
    )


def require_chat_vision_id(endpoint_id: str = "", legacy_model_id: str = "") -> str:
    v = resolve_chat_vision_id(endpoint_id or None, legacy_model_id or None)
    if not v:
        raise HTTPException(
            status_code=400,
            detail=(
                "缺少 Chat/Vision 接入点：请在侧栏填写 chat_vision_endpoint_id（ep-…），"
                "或在 backend/.env 配置 CHAT_VISION_ENDPOINT_ID / CHAT_VISION_MODEL_ID。"
            ),
        )
    return v


def require_embedding_id(endpoint_id: str = "", legacy_model_id: str = "") -> str:
    v = resolve_embedding_id(endpoint_id or None, legacy_model_id or None)
    if not v:
        raise HTTPException(
            status_code=400,
            detail=(
                "缺少 Embedding 接入点：请在侧栏填写 embedding_endpoint_id（ep-…），"
                "或在 backend/.env 配置 EMBEDDING_ENDPOINT_ID / EMBEDDING_MODEL_ID。"
            ),
        )
    return v


def resolve_api_key(request_value: str | None = None) -> str:
    return (request_value or "").strip() or (os.getenv("ARK_API_KEY") or "").strip()


def debug_log(run_id: str, hypothesis_id: str, location: str, message: str, data: dict[str, Any]) -> None:
    # #region agent log
    try:
        log_path = Path(__file__).resolve().parent.parent / ".cursor" / "debug-0a12b2.log"
        log_path.parent.mkdir(parents=True, exist_ok=True)
        payload = {
            "sessionId": "0a12b2",
            "runId": run_id,
            "hypothesisId": hypothesis_id,
            "location": location,
            "message": message,
            "data": data,
            "timestamp": int(datetime.now(timezone.utc).timestamp() * 1000),
        }
        with log_path.open("a", encoding="utf-8") as f:
            f.write(json.dumps(payload, ensure_ascii=False) + "\n")
    except Exception:
        pass
    # #endregion


@app.get("/api/config-check")
def config_check() -> dict[str, Any]:
    ark_api_key = (os.getenv("ARK_API_KEY") or "").strip()
    embedding_id = (
        (os.getenv("EMBEDDING_ENDPOINT_ID") or "").strip()
        or (os.getenv("EMBEDDING_MODEL_ID") or "").strip()
    )
    chat_id = (
        (os.getenv("CHAT_VISION_ENDPOINT_ID") or "").strip()
        or (os.getenv("CHAT_VISION_MODEL_ID") or "").strip()
    )

    def mask(value: str) -> str:
        if not value:
            return ""
        if len(value) <= 8:
            return "*" * len(value)
        return f"{value[:4]}***{value[-4:]}"

    return {
        "ok": bool(ark_api_key and embedding_id and chat_id),
        "vars": {
            "ARK_API_KEY": {
                "loaded": bool(ark_api_key),
                "masked": mask(ark_api_key),
            },
            "EMBEDDING_ENDPOINT_ID": {
                "loaded": bool((os.getenv("EMBEDDING_ENDPOINT_ID") or "").strip()),
                "value": (os.getenv("EMBEDDING_ENDPOINT_ID") or "").strip(),
            },
            "EMBEDDING_MODEL_ID": {
                "loaded": bool((os.getenv("EMBEDDING_MODEL_ID") or "").strip()),
                "value": (os.getenv("EMBEDDING_MODEL_ID") or "").strip(),
            },
            "CHAT_VISION_ENDPOINT_ID": {
                "loaded": bool((os.getenv("CHAT_VISION_ENDPOINT_ID") or "").strip()),
                "value": (os.getenv("CHAT_VISION_ENDPOINT_ID") or "").strip(),
            },
            "CHAT_VISION_MODEL_ID": {
                "loaded": bool((os.getenv("CHAT_VISION_MODEL_ID") or "").strip()),
                "value": (os.getenv("CHAT_VISION_MODEL_ID") or "").strip(),
            },
        },
    }


@app.get("/api/model-options")
def model_options() -> dict[str, Any]:
    """兼容旧前端；当前以侧栏手填 ep- 接入点为主，此处仅返回 .env 中的默认值（若有）。"""
    env_chat = resolve_chat_vision_id(None, None)
    env_embed = resolve_embedding_id(None, None)
    return {
        "chat_vision_models": [env_chat] if env_chat else [],
        "embedding_models": [env_embed] if env_embed else [],
        "defaults": {
            "chat_vision_endpoint_id": env_chat,
            "embedding_endpoint_id": env_embed,
            "chat_vision_model_id": env_chat,
            "embedding_model_id": env_embed,
        },
    }


@app.post("/api/model-probe")
def model_probe(payload: ModelProbeRequest) -> dict[str, Any]:
    run_id = f"probe-{int(datetime.now(timezone.utc).timestamp() * 1000)}"
    # #region agent log
    debug_log(
        run_id,
        "H1",
        "backend/main.py:model_probe:start",
        "model_probe request received",
        {
            "has_request_api_key": bool((payload.api_key or "").strip()),
            "request_chat_endpoint": (payload.chat_vision_endpoint_id or "").strip(),
            "request_embedding_endpoint": (payload.embedding_endpoint_id or "").strip(),
            "request_chat_legacy": (payload.chat_vision_model_id or "").strip(),
            "request_embedding_legacy": (payload.embedding_model_id or "").strip(),
        },
    )
    # #endregion
    api_key = resolve_api_key(payload.api_key)
    if not api_key:
        raise HTTPException(status_code=400, detail="缺少 API Key，请在齿轮填写或在 backend/.env 配置 ARK_API_KEY")

    chat_model_id = resolve_chat_vision_id(
        payload.chat_vision_endpoint_id,
        payload.chat_vision_model_id,
    )
    embedding_model_id = resolve_embedding_id(
        payload.embedding_endpoint_id,
        payload.embedding_model_id,
    )
    if not chat_model_id:
        raise HTTPException(
            status_code=400,
            detail="缺少 Chat/Vision 接入点 ID（请求体或 .env 均未配置）。",
        )
    if not embedding_model_id:
        raise HTTPException(
            status_code=400,
            detail="缺少 Embedding 接入点 ID（请求体或 .env 均未配置）。",
        )
    # #region agent log
    debug_log(
        run_id,
        "H2",
        "backend/main.py:model_probe:resolved-config",
        "resolved probe config",
        {
            "api_key_len": len(api_key),
            "api_key_from_request": bool((payload.api_key or "").strip()),
            "chat_model_id": chat_model_id,
            "embedding_model_id": embedding_model_id,
        },
    )
    # #endregion
    client = OpenAI(base_url=DOUBAO_BASE_URL, api_key=api_key)

    embedding_ok = False
    embedding_error = ""
    chat_ok = False
    chat_error = ""

    try:
        embed_texts(
            ["模型连通性测试"],
            embedding_model_id=embedding_model_id,
            api_key=api_key,
        )
        embedding_ok = True
    except Exception as exc:
        embedding_error = str(exc)
        app_logger.warning(
            "model_probe Embedding 探测失败（详情已写入响应 embedding.error）：%s",
            embedding_error[:2000],
        )
        # #region agent log
        debug_log(
            run_id,
            "H3",
            "backend/main.py:model_probe:embedding-error",
            "embedding probe failed",
            {
                "embedding_model_id": embedding_model_id,
                "error": embedding_error[:300],
            },
        )
        # #endregion

    try:
        client.chat.completions.create(
            model=chat_model_id,
            messages=[{"role": "user", "content": "请回复：ok"}],
            max_tokens=8,
            temperature=0,
        )
        chat_ok = True
    except Exception as exc:
        chat_error = str(exc)
        # #region agent log
        debug_log(
            run_id,
            "H4",
            "backend/main.py:model_probe:chat-error",
            "chat probe failed",
            {
                "chat_model_id": chat_model_id,
                "error": chat_error[:300],
            },
        )
        # #endregion

    # #region agent log
    debug_log(
        run_id,
        "H5",
        "backend/main.py:model_probe:result",
        "model_probe finished",
        {
            "ok": embedding_ok and chat_ok,
            "embedding_ok": embedding_ok,
            "chat_ok": chat_ok,
        },
    )
    # #endregion
    return {
        "ok": embedding_ok and chat_ok,
        "embedding": {
            "model_id": embedding_model_id,
            "ok": embedding_ok,
            "error": embedding_error,
        },
        "chat_vision": {
            "model_id": chat_model_id,
            "ok": chat_ok,
            "error": chat_error,
        },
    }


@app.on_event("startup")
def on_startup() -> None:
    init_db()


@app.get("/api/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/api/projects", response_model=ProjectRead)
def create_project(payload: ProjectCreate) -> Project:
    name = payload.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="project name is required")
    with get_session() as session:
        project = Project(name=name)
        session.add(project)
        session.commit()
        session.refresh(project)
        ensure_project_collection(project.id or 0)
        project_dir_by_id(project.id or 0).mkdir(parents=True, exist_ok=True)
        return project


@app.get("/api/projects", response_model=list[ProjectRead])
def list_projects() -> list[Project]:
    with get_session() as session:
        statement = select(Project).order_by(Project.created_at.desc())
        return list(session.exec(statement))


@app.get("/api/projects/{project_id}/state", response_model=ProjectStateRead)
def get_project_state(project_id: int) -> ProjectStateRead:
    get_project_or_404(project_id)
    with get_session() as session:
        doc_statement = select(ProjectDocument).where(ProjectDocument.project_id == project_id)
        project_doc = session.exec(doc_statement).first()
        msg_statement = (
            select(ChatSessionMessage)
            .where(ChatSessionMessage.project_id == project_id)
            .order_by(ChatSessionMessage.created_at.asc(), ChatSessionMessage.id.asc())
        )
        messages = list(session.exec(msg_statement))
    message_payload = [
        ChatMessageRead(
            id=item.id or 0,
            project_id=item.project_id,
            role=item.role,
            content=item.content,
            created_at=item.created_at,
        )
        for item in messages
    ]
    return ProjectStateRead(
        project_id=project_id,
        document=project_doc.content if project_doc else "",
        messages=message_payload,
    )


@app.put("/api/projects/{project_id}/document", response_model=ProjectStateRead)
def save_project_document(project_id: int, payload: ProjectDocumentUpdate) -> ProjectStateRead:
    get_project_or_404(project_id)
    with get_session() as session:
        doc_statement = select(ProjectDocument).where(ProjectDocument.project_id == project_id)
        project_doc = session.exec(doc_statement).first()
        if not project_doc:
            project_doc = ProjectDocument(project_id=project_id, content=payload.content)
        else:
            project_doc.content = payload.content
            project_doc.updated_at = datetime.now(timezone.utc)
        session.add(project_doc)
        session.commit()
    return get_project_state(project_id)


def _sanitize_markdown_for_word_export(md: str) -> str:
    """去掉 [doc:…] 依据占位与 BlockNote 等导出的「单行反斜杠」硬换行占位。"""
    text0 = (md or "").strip()
    text0 = re.sub(r"\[doc\s*:\s*[^\]]+\]", "", text0, flags=re.IGNORECASE)
    lines: list[str] = []
    for line in text0.splitlines():
        s = line.strip()
        if s == "\\" or s == "\\\\":
            continue
        lines.append(line)
    text = "\n".join(lines)
    while "\n\n\n" in text:
        text = text.replace("\n\n\n", "\n\n")
    return text.strip()


def _markdown_to_docx_bytes(md_text: str) -> bytes:
    from io import BytesIO

    import markdown as md_lib
    from docx import Document
    from htmldocx import HtmlToDocx

    text = _sanitize_markdown_for_word_export(md_text or "")
    doc = Document()
    if not text:
        doc.add_paragraph("（空文档）")
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # 不用 nl2br：与 BlockNote 已有段落/换行叠加后易产生异常断行；extra 已含表格等
    html = md_lib.markdown(
        text,
        extensions=["extra", "sane_lists"],
    )
    wrapped = f"<div>{html}</div>"
    try:
        HtmlToDocx().add_html_to_document(wrapped, doc)
    except Exception as exc:
        app_logger.warning("htmldocx failed, using plain paragraphs: %s", exc)
        doc = Document()
        for block in text.split("\n\n"):
            doc.add_paragraph(block.replace("\n", " "))
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.post("/api/projects/{project_id}/export/docx")
def export_project_docx(project_id: int, payload: DocumentExportRequest) -> Response:
    project = get_project_or_404(project_id)
    md_body = (payload.content or "").strip()
    if not md_body:
        with get_session() as session:
            doc_statement = select(ProjectDocument).where(ProjectDocument.project_id == project_id)
            project_doc = session.exec(doc_statement).first()
            md_body = (project_doc.content if project_doc else "").strip()
    data = _markdown_to_docx_bytes(md_body)
    safe_slug = re.sub(r"[^\w\u4e00-\u9fff\-]+", "_", (project.name or str(project_id)).strip())[:80]
    filename_ascii = f"{safe_slug or project_id}.docx"
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename_ascii}"',
        },
    )


@app.post("/api/projects/{project_id}/messages", response_model=ChatMessageRead)
def append_project_message(project_id: int, payload: ChatMessageCreate) -> ChatMessageRead:
    get_project_or_404(project_id)
    role = payload.role.strip().lower()
    if role not in {"user", "assistant"}:
        raise HTTPException(status_code=400, detail="role must be user or assistant")
    with get_session() as session:
        message = ChatSessionMessage(project_id=project_id, role=role, content=payload.content)
        session.add(message)
        session.commit()
        session.refresh(message)
        return ChatMessageRead(
            id=message.id or 0,
            project_id=message.project_id,
            role=message.role,
            content=message.content,
            created_at=message.created_at,
        )


@app.delete("/api/projects/{project_id}/messages")
def clear_project_messages(project_id: int) -> dict[str, int]:
    get_project_or_404(project_id)
    with get_session() as session:
        statement = select(ChatSessionMessage).where(ChatSessionMessage.project_id == project_id)
        records = list(session.exec(statement))
        deleted = len(records)
        for row in records:
            session.delete(row)
        session.commit()
    return {"project_id": project_id, "deleted": deleted}


@app.patch("/api/projects/{project_id}", response_model=ProjectRead)
def rename_project(project_id: int, payload: ProjectUpdate) -> Project:
    name = payload.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="project name is required")
    with get_session() as session:
        project = session.get(Project, project_id)
        if not project:
            raise HTTPException(status_code=404, detail="project not found")
        project.name = name
        session.add(project)
        session.commit()
        session.refresh(project)
        return project


@app.delete("/api/projects/{project_id}")
def delete_project(project_id: int) -> dict[str, Any]:
    with get_session() as session:
        project = session.get(Project, project_id)
        if not project:
            raise HTTPException(status_code=404, detail="project not found")

        doc_statement = select(ProjectDocument).where(ProjectDocument.project_id == project_id)
        project_doc = session.exec(doc_statement).first()
        if project_doc:
            session.delete(project_doc)

        msg_statement = select(ChatSessionMessage).where(ChatSessionMessage.project_id == project_id)
        records = list(session.exec(msg_statement))
        for row in records:
            session.delete(row)

        session.delete(project)
        session.commit()

    project_dir = project_dir_by_id(project_id)
    if project_dir.exists():
        shutil.rmtree(project_dir, ignore_errors=True)
    delete_project_collection(project_id)
    return {"project_id": project_id, "status": "deleted"}

@app.post("/api/evidence/delete")
def delete_evidence(payload: EvidenceDeleteRequest) -> dict[str, Any]:
    project_id = int(payload.project_id)
    get_project_or_404(project_id)
    project_dir = project_dir_by_id(project_id)
    project_dir.mkdir(parents=True, exist_ok=True)

    # Resolve credentials (request first)
    api_key = resolve_api_key(payload.api_key)
    chat_vision_id = resolve_chat_vision_id(payload.chat_vision_endpoint_id, None)
    embedding_id = resolve_embedding_id(payload.embedding_endpoint_id, payload.embedding_model_id) or ""
    if not api_key:
        raise HTTPException(status_code=400, detail="缺少 API Key")
    if not embedding_id:
        raise HTTPException(status_code=400, detail="缺少 Embedding 接入点 ID")

    full = _resolve_project_file_or_400(project_id, payload.rel_path)
    try:
        full.unlink(missing_ok=True)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"删除失败：{exc}") from exc

    # Re-embed all remaining evidence for consistency
    try:
        delete_project_collection(project_id)
        ensure_project_collection(project_id)
        stats = upsert_project_evidence(
            project_id=project_id,
            chat_vision_endpoint_id=chat_vision_id,
            embedding_model_id=embedding_id,
            api_key=api_key,
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"向量重建失败：{exc}") from exc

    files = list_project_files(project_dir)
    return {
        "ok": True,
        "project_id": str(project_id),
        "deleted_rel_path": payload.rel_path,
        "files": files,
        "total": len(files),
        "vector_upsert": {"upserted": stats.get("upserted", 0), "files": stats.get("files", 0)},
    }


@app.post("/api/upload")
async def upload_file(
    project_id: int = Form(...),
    api_key: str | None = Form(default=None),
    chat_vision_endpoint_id: str | None = Form(default=None),
    embedding_endpoint_id: str | None = Form(default=None),
    embedding_model_id: str | None = Form(default=None),
    file: UploadFile = File(...),
) -> dict[str, object]:
    filename = file.filename or ""
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="unsupported file type")

    project_dir = project_dir_by_id(project_id)
    project_dir.mkdir(parents=True, exist_ok=True)

    safe_name = Path(filename).name
    uploaded_path = project_dir / safe_name
    save_upload_file(file, uploaded_path)

    if suffix == ".zip":
        try:
            extract_zip_safely(uploaded_path, project_dir)
            uploaded_path.unlink(missing_ok=True)
        except zipfile.BadZipFile as exc:
            uploaded_path.unlink(missing_ok=True)
            raise HTTPException(status_code=400, detail="invalid zip file") from exc
        except Exception as exc:
            uploaded_path.unlink(missing_ok=True)
            raise HTTPException(status_code=400, detail=f"zip extract failed: {exc}") from exc

    for zip_file in project_dir.rglob("*.zip"):
        zip_file.unlink(missing_ok=True)

    normalize_xls_files_under(project_dir)

    upsert_stats: dict[str, Any] = {"upserted": 0, "files": 0, "skipped": True}
    effective_api_key = resolve_api_key(api_key)
    effective_chat_vision_endpoint_id = resolve_chat_vision_id(chat_vision_endpoint_id, None)
    effective_embedding_model_id = resolve_embedding_id(embedding_endpoint_id, embedding_model_id) or None
    if effective_api_key:
        try:
            ensure_project_collection(project_id)
            stats = upsert_project_evidence(
                project_id=project_id,
                embedding_model_id=effective_embedding_model_id,
                api_key=effective_api_key,
                chat_vision_endpoint_id=effective_chat_vision_endpoint_id,
            )
            upsert_stats = {"upserted": stats["upserted"], "files": stats["files"], "skipped": False}
        except Exception as exc:
            upsert_stats = {"upserted": 0, "files": 0, "skipped": False, "error": str(exc)}

    files = list_project_files(project_dir)
    return {
        "project_id": str(project_id),
        "files": files,
        "total": len(files),
        "vector_upsert": upsert_stats,
    }


@app.post("/api/upload-template")
async def upload_template(
    project_id: int = Form(...),
    file: UploadFile = File(...),
) -> dict[str, str]:
    filename = file.filename or ""
    suffix = Path(filename).suffix.lower()
    if suffix not in TEMPLATE_ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="template only supports PDF/Word")

    project_dir = project_dir_by_id(project_id)
    template_dir = project_dir / "template"
    template_dir.mkdir(parents=True, exist_ok=True)

    for old in template_dir.iterdir():
        if old.is_file():
            old.unlink(missing_ok=True)

    safe_name = Path(filename).name
    template_path = template_dir / safe_name
    save_upload_file(file, template_path)
    return {
        "project_id": str(project_id),
        "template_name": safe_name,
        "template_type": detect_file_type(template_path),
        "status": "ok",
    }

@app.get("/api/projects/{project_id}/files")
def list_files(project_id: int) -> dict[str, Any]:
    get_project_or_404(project_id)
    project_dir = project_dir_by_id(project_id)
    project_dir.mkdir(parents=True, exist_ok=True)
    files = list_project_files(project_dir)
    return {"project_id": str(project_id), "files": files, "total": len(files)}


@app.get("/api/projects/{project_id}/evidence/path")
def evidence_path(project_id: int, rel_path: str) -> dict[str, str]:
    get_project_or_404(project_id)
    full = _resolve_project_file_or_400(project_id, rel_path)
    return {"project_id": str(project_id), "rel_path": rel_path, "abs_path": str(full)}


@app.get("/api/projects/{project_id}/evidence/preview")
def evidence_preview(project_id: int, rel_path: str) -> object:
    get_project_or_404(project_id)
    full = _resolve_project_file_or_400(project_id, rel_path)
    ext = full.suffix.lower()
    if ext in IMAGE_EXTENSIONS:
        media_type, _ = mimetypes.guess_type(str(full))
        return FileResponse(str(full), media_type=media_type or "image/jpeg", headers={"Content-Disposition": "inline"})
    if ext == ".pdf":
        return FileResponse(str(full), media_type="application/pdf", headers={"Content-Disposition": "inline"})
    if ext in {".txt", ".md"}:
        try:
            text = full.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            text = ""
        return PlainTextResponse(text)
    if ext in {".docx"}:
        return PlainTextResponse(extract_docx_text(full))
    if ext == ".xlsx":
        return PlainTextResponse(extract_xlsx_text(full))
    if ext == ".csv":
        return PlainTextResponse(extract_csv_text(full))
    # fallback: try plain text extraction
    return PlainTextResponse(extract_text_by_file(full))


@app.get("/api/projects/{project_id}/evidence/chunks")
def evidence_chunks(project_id: int, rel_path: str, max_chunks: int = 60) -> dict[str, Any]:
    get_project_or_404(project_id)
    full = _resolve_project_file_or_400(project_id, rel_path)
    text = trim_text(extract_text_by_file(full))
    chunks = chunk_text(text)
    max_n = max(1, min(int(max_chunks), 200))
    truncated = len(chunks) > max_n
    return {
        "project_id": str(project_id),
        "rel_path": rel_path,
        "chunks": chunks[:max_n],
        "total_chunks": len(chunks),
        "truncated": truncated,
    }


@app.get("/api/projects/{project_id}/evidence/keyfacts")
def evidence_keyfacts(project_id: int, rel_path: str) -> dict[str, Any]:
    """
    读取某个证据文件对应的 KeyFacts 缓存（由 OCR 流程同步生成/重扫更新）。
    若缓存不存在/指纹不一致/内容为空，则返回 missing，由前端提示用户先重扫 OCR。
    """
    get_project_or_404(project_id)
    full = _resolve_project_file_or_400(project_id, rel_path)
    project_dir = project_dir_by_id(project_id)

    fingerprint = compute_file_fingerprint(full)
    cache_file = _keyfacts_cache_file(project_dir, full)
    cached = _read_keyfacts_cache(cache_file, fingerprint)
    if cached:
        full_text = cached.get("keyfacts_full") or ""
        compact_text = cached.get("keyfacts_compact") or ""
        if full_text.strip():
            return {
                "status": "ok",
                "project_id": str(project_id),
                "rel_path": rel_path,
                "keyfacts_full": full_text,
                "keyfacts_compact": compact_text,
            }

    return {
        "status": "missing",
        "project_id": str(project_id),
        "rel_path": rel_path,
        "message": "未生成 KeyFacts，请先使用“重新扫描 OCR 并重建向量”。",
    }


def _ocr_rebuild_prepare(project_id: int, payload: OcrRebuildRequest) -> tuple[Path, list[Path], str, str, str, str]:
    get_project_or_404(project_id)
    api_key = resolve_api_key(payload.api_key)
    if not api_key:
        raise HTTPException(status_code=400, detail="缺少 API Key")

    chat_vision_id = resolve_chat_vision_id(payload.chat_vision_endpoint_id, None)
    if not chat_vision_id:
        raise HTTPException(
            status_code=400,
            detail="缺少 chat_vision_endpoint_id（OCR 需要）。请侧栏填写或在 backend/.env 配置。",
        )

    embedding_id = resolve_embedding_id(payload.embedding_endpoint_id, payload.embedding_model_id) or ""
    if not embedding_id and not os.getenv("EMBEDDING_ENDPOINT_ID") and not os.getenv("EMBEDDING_MODEL_ID"):
        raise HTTPException(status_code=400, detail="缺少 Embedding 接入点 ID（用于重建向量库）。")

    project_dir = project_dir_by_id(project_id)
    if not project_dir.exists():
        raise HTTPException(status_code=404, detail="project data not found")

    scope_rel_path = (payload.rel_path or "").strip()
    if scope_rel_path:
        file_path = _resolve_project_file_or_400(project_id, scope_rel_path)
        target_files = [file_path]
    else:
        normalize_xls_files_under(project_dir)
        target_files = list(iter_evidence_files(project_dir))

    return project_dir, target_files, api_key, chat_vision_id, embedding_id, scope_rel_path


def _iter_ocr_rebuild_steps(
    project_id: int,
    project_dir: Path,
    target_files: list[Path],
    api_key: str,
    chat_vision_id: str,
    embedding_id: str,
    scope_rel_path: str,
) -> Iterable[dict[str, Any]]:
    n = len(target_files)
    total_steps = max(1, n + 1)
    yield {
        "kind": "progress",
        "phase": "start",
        "current": 0,
        "total": total_steps,
        "message": f"共 {n} 个证据文件，开始 OCR / KeyFacts…",
    }

    ocr_attempted = 0
    ocr_updated = 0
    keyfacts_attempted = 0
    keyfacts_updated = 0

    for idx, file_path in enumerate(target_files):
        yield {
            "kind": "progress",
            "phase": "file",
            "current": idx + 1,
            "total": total_steps,
            "message": f"（{idx + 1}/{n}）{file_path.name}",
        }
        ext = file_path.suffix.lower()
        is_media = ext in IMAGE_EXTENSIONS or ext == ".pdf"

        ocr_text: str = ""
        if is_media:
            ocr_attempted += 1
            before = extract_text_by_file(file_path)
            ocr_text = get_or_create_ocr_text(
                project_dir=project_dir,
                file_path=file_path,
                api_key=api_key,
                chat_vision_endpoint_id=chat_vision_id,
                reocr_if_empty=True,
            )
            if ocr_text and ocr_text.strip() and (not before or not before.strip()):
                ocr_updated += 1

        source_text = ocr_text if is_media else extract_text_by_file(file_path)
        fingerprint = compute_file_fingerprint(file_path)
        cache_file = _keyfacts_cache_file(project_dir, file_path)
        cached_before = _read_keyfacts_cache(cache_file, fingerprint)
        cached_full = str((cached_before or {}).get("keyfacts_full") or "").strip()

        keyfacts_attempted += 1
        need_keyfacts = not cached_full
        if need_keyfacts:
            keyfacts = get_or_create_keyfacts_text(
                project_dir=project_dir,
                file_path=file_path,
                api_key=api_key,
                chat_vision_endpoint_id=chat_vision_id,
                ocr_text=source_text,
                fingerprint=fingerprint,
                rekeyfacts_if_empty=True,
                force=False,
            )
            if str(keyfacts.get("keyfacts_full") or "").strip():
                keyfacts_updated += 1

    yield {
        "kind": "progress",
        "phase": "vector",
        "current": total_steps,
        "total": total_steps,
        "message": "正在向量化并写入检索库…",
    }
    delete_project_collection(project_id)
    ensure_project_collection(project_id)
    stats = upsert_project_evidence(
        project_id=project_id,
        chat_vision_endpoint_id=chat_vision_id,
        embedding_model_id=embedding_id or None,
        api_key=api_key,
    )

    files = list_project_files(project_dir)
    yield {
        "kind": "done",
        "result": {
            "ok": True,
            "project_id": str(project_id),
            "scope_rel_path": scope_rel_path,
            "ocr_attempted": ocr_attempted,
            "ocr_updated": ocr_updated,
            "keyfacts_attempted": keyfacts_attempted,
            "keyfacts_updated": keyfacts_updated,
            "vector_upsert": stats,
            "files": files,
        },
    }


@app.post("/api/projects/{project_id}/ocr/rebuild", response_model=None)
def ocr_rebuild_endpoint(project_id: int, payload: OcrRebuildRequest) -> dict[str, Any] | StreamingResponse:
    """
    确保 `.ocr_cache`（图片/PDF）在缺失或缓存为空时补齐；对所有证据文件补齐/更新 KeyFacts（含压缩版）；最后重建向量库。
    payload.stream=true 时返回 SSE：progress → done（body 字段与原先 JSON 一致）。
    """
    project_dir, target_files, api_key, chat_vision_id, embedding_id, scope_rel_path = _ocr_rebuild_prepare(
        project_id, payload
    )

    if payload.stream:

        def sse_gen() -> Iterable[str]:
            try:
                for step in _iter_ocr_rebuild_steps(
                    project_id,
                    project_dir,
                    target_files,
                    api_key,
                    chat_vision_id,
                    embedding_id,
                    scope_rel_path,
                ):
                    if step["kind"] == "progress":
                        yield sse_event(
                            {
                                "type": "progress",
                                "phase": step.get("phase", ""),
                                "current": step.get("current", 0),
                                "total": step.get("total", 1),
                                "message": step.get("message", ""),
                            }
                        )
                    elif step["kind"] == "done":
                        yield sse_event({"type": "done", **step["result"]})
            except HTTPException as exc:
                detail = exc.detail if isinstance(exc.detail, str) else str(exc.detail)
                yield sse_event({"type": "error", "message": detail})
            except Exception as exc:
                yield sse_event({"type": "error", "message": str(exc)})

        return StreamingResponse(sse_gen(), media_type="text/event-stream")

    steps = list(
        _iter_ocr_rebuild_steps(
            project_id,
            project_dir,
            target_files,
            api_key,
            chat_vision_id,
            embedding_id,
            scope_rel_path,
        )
    )
    if not steps or steps[-1].get("kind") != "done":
        raise HTTPException(status_code=500, detail="OCR 重建未完成")
    return steps[-1]["result"]


def _build_legal_opinion_system_prompt(template_text: str) -> str:
    return f"""你是一个资深非诉律师。请根据提供的证据资料与用户指令，输出法律相关 Markdown 正文（可以是全新撰写，或在用户给出的「当前全文 Markdown」基础上改写、扩写或重排）。

输出要求（必须遵守）：
1. 只输出 Markdown 正文，不要输出任何额外说明、JSON、或代码块外壳（不要使用 ```）。
2. 标题层级使用 Markdown 标题语法：一级 `#`、二级 `##`、三级 `###`。
3. 列表使用 `-` 或 `1.`。
4. 表格必须使用标准 Markdown 表格语法，并包含表头与分隔行：
   | 列名1 | 列名2 |
   | --- | --- |
   | ... | ... |
   表格内不要设置「来源」「出处」「依据」等列；表格仅承载序号、案号、裁判结果、履行状态等业务信息。
5. 事实陈述须在合适位置用严格格式标注来源：[doc:文件名]，其中 `文件名` 必须与证据条目中的 `source_file` 一致。来源标注应放在**表格外**（如表格下方按行说明、或段末），不要单独占一列表格列。
6. 避免多余空行：段落之间通常最多空一行。

参考模板（仅作为内容组织与风格参考，不要原样复制模板的非正文说明部分）：
{template_text}"""


def iter_legal_opinion_generation_sse(
    *,
    project_dir: Path,
    api_key: str,
    chat_id: str,
    user_instruction_text: str,
    extra_history: list[dict[str, str]] | None = None,
) -> Iterable[str]:
    """模板 + 证据（原文/压缩档自动切换）+ 流式 delta；供 /api/generate 与 chat full_rewrite 共用。"""
    template_text = load_template_text(project_dir)
    system_prompt = _build_legal_opinion_system_prompt(template_text)
    max_context = _env_int("CHAT_MAX_CONTEXT_TOKENS", 256_000)
    reserve_out = _env_int("RESERVE_OUTPUT_TOKENS", 8_000)
    safety_margin = _env_int("SAFETY_MARGIN_TOKENS", 4_000)

    evidence_full = build_evidence_content_items(project_dir)
    est_full = estimate_prompt_tokens(
        system_prompt=system_prompt,
        instruction=user_instruction_text,
        evidence_items=evidence_full,
    )
    threshold = max(8_000, max_context - reserve_out - safety_margin)
    use_compact = est_full > threshold

    evidence_content = build_evidence_content_items_compact(project_dir) if use_compact else evidence_full
    est_used = (
        estimate_prompt_tokens(
            system_prompt=system_prompt,
            instruction=user_instruction_text,
            evidence_items=evidence_content,
        )
        if use_compact
        else est_full
    )
    mode = "压缩档" if use_compact else "原文档"

    user_content: list[dict[str, Any]] = [{"type": "text", "text": user_instruction_text}, *evidence_content]

    try:
        yield sse_event({"type": "hint", "message": f"本次生成使用：{mode}（预计输入≈{est_used} tokens，上限≈{max_context}）"})
        yield sse_event(
            {
                "type": "progress",
                "phase": "model",
                "indeterminate": True,
                "message": "正在连接模型并等待首段输出…",
            }
        )
        client = OpenAI(base_url=DOUBAO_BASE_URL, api_key=api_key)
        messages: list[dict[str, Any]] = [{"role": "system", "content": system_prompt}]
        if extra_history:
            messages.extend(extra_history)
        messages.append({"role": "user", "content": user_content})
        stream = client.chat.completions.create(
            model=chat_id,
            messages=messages,
            stream=True,
            temperature=0.2,
        )
        for chunk in stream:
            delta = ""
            if chunk.choices and chunk.choices[0].delta:
                delta = chunk.choices[0].delta.content or ""
            if delta:
                yield sse_event({"type": "delta", "content": delta})
        yield sse_event({"type": "done", "sources": []})
    except Exception as exc:
        yield sse_event({"type": "error", "message": str(exc)})


@app.post("/api/generate")
async def generate_legal_opinion(payload: GenerateRequest) -> StreamingResponse:
    api_key = resolve_api_key(payload.api_key)
    if not api_key:
        raise HTTPException(status_code=400, detail="缺少 API Key，请在齿轮填写或在 backend/.env 配置 ARK_API_KEY")

    project_dir = project_dir_by_id(payload.project_id)
    if not project_dir.exists():
        raise HTTPException(status_code=404, detail="project data not found")

    instruction = payload.instruction.strip() or "请生成法律意见书正文。"
    chat_id = require_chat_vision_id(
        payload.chat_vision_endpoint_id,
        payload.chat_vision_model_id,
    )
    return StreamingResponse(
        iter_legal_opinion_generation_sse(
            project_dir=project_dir,
            api_key=api_key,
            chat_id=chat_id,
            user_instruction_text=instruction,
            extra_history=None,
        ),
        media_type="text/event-stream",
    )


@app.post("/api/rag-query", response_model=RagQueryResponse)
async def rag_query(payload: RagQueryRequest) -> RagQueryResponse:
    query = payload.query.strip()
    api_key = resolve_api_key(payload.api_key)
    if not query:
        raise HTTPException(status_code=400, detail="query is required")
    if not api_key:
        raise HTTPException(status_code=400, detail="缺少 API Key，请在齿轮填写或在 backend/.env 配置 ARK_API_KEY")

    project_dir = project_dir_by_id(payload.project_id)
    if not project_dir.exists():
        raise HTTPException(status_code=404, detail="project data not found")

    meta_where = build_metadata_where(
        filter_doc_type=payload.filter_doc_type or None,
        filter_source_file=payload.filter_source_file or None,
    )
    embed_id = require_embedding_id(
        payload.embedding_endpoint_id,
        payload.embedding_model_id,
    )
    chat_id = require_chat_vision_id(
        payload.chat_vision_endpoint_id,
        payload.chat_vision_model_id,
    )
    try:
        retrieval = query_project_chunks(
            project_id=payload.project_id,
            query=query,
            top_k=payload.top_k,
            embedding_model_id=embed_id,
            where=meta_where,
            api_key=api_key,
        )
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"RAG 检索失败：{exc}") from exc
    docs = retrieval.get("documents", [])
    metas = retrieval.get("metadatas", [])
    if not docs:
        return RagQueryResponse(answer="当前项目还没有可检索内容，请先上传并完成向量入库。", sources=[])

    context_lines: list[str] = []
    source_names: list[str] = []
    for idx, doc in enumerate(docs):
        meta = metas[idx] if idx < len(metas) else {}
        doc_name = str(meta.get("source_file") or meta.get("doc_name", "未知文件"))
        source_names.append(doc_name)
        context_lines.append(f"[doc:{doc_name}] {doc}")
    context = "\n\n".join(context_lines)

    system_prompt = (
        "你是资深非诉律师的项目检索助手。"
        "请仅基于给出的检索片段回答用户问题，不能编造事实。"
        "回答要简洁专业，并在关键句末用 [doc:文件名] 标注来源。"
    )
    user_prompt = f"用户问题：{query}\n\n检索片段：\n{context}"

    try:
        client = OpenAI(base_url=DOUBAO_BASE_URL, api_key=api_key)
        completion = client.chat.completions.create(
            model=chat_id,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
        )
        answer = (completion.choices[0].message.content or "").strip()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"rag query failed: {exc}") from exc

    unique_sources: list[str] = []
    for name in source_names:
        if name not in unique_sources:
            unique_sources.append(name)
    return RagQueryResponse(answer=answer or "未生成有效回答。", sources=unique_sources)


@app.post("/api/rag-query/stream")
async def rag_query_stream(payload: RagQueryRequest) -> StreamingResponse:
    return StreamingResponse(iter_rag_query_sse(payload), media_type="text/event-stream")


def _rag_payload_from_project_chat(project_id: int, p: ProjectChatStreamRequest) -> RagQueryRequest:
    return RagQueryRequest(
        project_id=project_id,
        api_key=p.api_key,
        chat_vision_endpoint_id=p.chat_vision_endpoint_id,
        embedding_endpoint_id=p.embedding_endpoint_id,
        chat_vision_model_id=p.chat_vision_model_id,
        embedding_model_id=p.embedding_model_id,
        query=p.message,
        top_k=p.top_k,
        filter_doc_type=p.filter_doc_type,
        filter_source_file=p.filter_source_file,
    )


def iter_project_chat_sse(project_id: int, payload: ProjectChatStreamRequest) -> Iterable[str]:
    message = (payload.message or "").strip()
    if not message:
        yield sse_event({"type": "error", "message": "message is required"})
        yield sse_event({"type": "done", "sources": []})
        return

    api_key = resolve_api_key(payload.api_key)
    if not api_key:
        yield sse_event({"type": "error", "message": "缺少 API Key，请在齿轮填写或在 backend/.env 配置 ARK_API_KEY"})
        yield sse_event({"type": "done", "sources": []})
        return

    project_dir = project_dir_by_id(project_id)
    if not project_dir.exists():
        yield sse_event({"type": "error", "message": "project data not found"})
        yield sse_event({"type": "done", "sources": []})
        return

    mode = payload.mode
    if mode == ChatStreamMode.search:
        yield from iter_rag_query_sse(_rag_payload_from_project_chat(project_id, payload))
        return

    if mode == ChatStreamMode.partial_rewrite:
        sel = payload.selection
        if not sel or not str(sel.text).strip():
            yield sse_event({"type": "error", "message": "局部改写需要先在编辑器中引用选区"})
            yield sse_event({"type": "done", "sources": []})
            return

    try:
        chat_id = require_chat_vision_id(
            payload.chat_vision_endpoint_id,
            payload.chat_vision_model_id,
        )
    except HTTPException as exc:
        detail = exc.detail
        msg = detail if isinstance(detail, str) else str(detail)
        yield sse_event({"type": "error", "message": msg})
        yield sse_event({"type": "done", "sources": []})
        return

    hist = _chat_history_openai_messages(payload.history)
    client = OpenAI(base_url=DOUBAO_BASE_URL, api_key=api_key)

    if mode == ChatStreamMode.qa:
        system_prompt = (
            "你是专业法律顾问。请结合多轮对话回答用户问题，语气专业、简洁。"
            "不要编造具体案例或法条条文；若不确定请明确说明。"
        )
        messages: list[dict[str, str]] = [{"role": "system", "content": system_prompt}, *hist, {"role": "user", "content": message}]

        try:
            stream = client.chat.completions.create(
                model=chat_id,
                messages=messages,
                stream=True,
                temperature=0.45,
            )
            for chunk in stream:
                delta = ""
                if chunk.choices and chunk.choices[0].delta:
                    delta = chunk.choices[0].delta.content or ""
                if delta:
                    yield sse_event({"type": "delta", "content": delta})
            yield sse_event({"type": "done", "sources": []})
        except Exception as exc:
            yield sse_event({"type": "error", "message": str(exc)})
        return

    if mode == ChatStreamMode.full_rewrite:
        doc = _truncate_for_chat_context(payload.document_markdown or "")
        instruction_block = f"【用户指令】\n{message}\n\n【当前全文 Markdown】\n{doc}"
        yield from iter_legal_opinion_generation_sse(
            project_dir=project_dir,
            api_key=api_key,
            chat_id=chat_id,
            user_instruction_text=instruction_block,
            extra_history=hist,
        )
        return

    if mode == ChatStreamMode.partial_rewrite:
        sel_text = (payload.selection.text or "").strip()
        label = (payload.selection.label or "").strip() or "（未标注位置）"
        full_ctx = _truncate_for_chat_context(payload.document_markdown or "", max_chars=min(CHAT_DOCUMENT_MARKDOWN_MAX, 80_000))
        system_prompt = (
            "你是法律文档编辑。用户给出了文稿中的「待改片段」以及全文上下文（可能截断）。"
            "你只能改写该片段的表述以符合用户指令，须与上下文在事实与语气上衔接；不得输出片段之外的内容。"
            "只输出改写后的片段正文（与选区相同的 Markdown 粒度即可），不要任何开场白、标题或整篇文档。"
        )
        user_body = (
            f"【用户指令】\n{message}\n\n"
            f"【选区位置说明】\n{label}\n\n"
            f"【待改选区】\n{sel_text}\n\n"
            f"【全文上下文（可能截断）】\n{full_ctx if full_ctx.strip() else '（未提供）'}"
        )
        messages = [{"role": "system", "content": system_prompt}, *hist, {"role": "user", "content": user_body}]
        try:
            stream = client.chat.completions.create(
                model=chat_id,
                messages=messages,
                stream=True,
                temperature=0.15,
            )
            for chunk in stream:
                delta = ""
                if chunk.choices and chunk.choices[0].delta:
                    delta = chunk.choices[0].delta.content or ""
                if delta:
                    yield sse_event({"type": "delta", "content": delta})
            yield sse_event({"type": "done", "sources": []})
        except Exception as exc:
            yield sse_event({"type": "error", "message": str(exc)})
        return

    yield sse_event({"type": "error", "message": f"unknown chat mode: {mode}"})
    yield sse_event({"type": "done", "sources": []})


@app.post("/api/projects/{project_id}/chat/stream")
async def project_chat_stream(project_id: int, payload: ProjectChatStreamRequest) -> StreamingResponse:
    return StreamingResponse(iter_project_chat_sse(project_id, payload), media_type="text/event-stream")
